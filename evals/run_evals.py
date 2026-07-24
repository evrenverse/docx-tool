"""Reproducible agent workflow eval for docx-tool."""

from __future__ import annotations

import json
import shutil
import subprocess
import tempfile
import time
from pathlib import Path

from docx import Document


def run(binary: str, *args: str) -> str:
    result = subprocess.run(
        [binary, *args],
        check=False,
        capture_output=True,
        text=True,
        timeout=60,
    )
    if result.returncode != 0:
        raise RuntimeError(
            f"{' '.join(args)} exited {result.returncode}: {result.stderr or result.stdout}"
        )
    return result.stdout


def create_fixture(path: Path) -> None:
    document = Document()
    document.add_heading("Status report", level=1)
    document.add_paragraph("Owner: Example Team")
    document.add_paragraph("Total: 1000")
    document.save(path)


def main() -> None:
    started = time.monotonic()
    binary = shutil.which("docx-tool")
    if binary is None:
        raise SystemExit("docx-tool not found on PATH; run with `uv run python evals/run_evals.py`")

    with tempfile.TemporaryDirectory(prefix="docx-tool-eval-") as temporary:
        root = Path(temporary)
        source = root / "report.docx"
        output = root / "updated.docx"
        changes = root / "changes.json"
        create_fixture(source)
        changes.write_text(
            json.dumps([{"type": "replace", "search": "Example Team", "value": "Agent Team"}]),
            encoding="utf-8",
        )

        run(binary, "info", str(source), "--json")
        found = json.loads(run(binary, "find", str(source), "Owner", "--json"))
        if found["total"] != 1:
            raise RuntimeError(f"expected one owner match, got {found['total']}")
        run(binary, "read", str(source), "--paragraphs", "2", "--json")
        run(
            binary,
            "write",
            str(source),
            str(changes),
            "--output",
            str(output),
            "--json",
        )
        final = json.loads(run(binary, "read", str(output), "--paragraphs", "2", "--json"))
        if final["paragraphs"][0]["text"] != "Owner: Agent Team":
            raise RuntimeError(f"unexpected final paragraph: {final['paragraphs']}")

    print(
        json.dumps(
            {
                "schema_version": "1",
                "tool": "docx-tool",
                "eval": "locate-edit-verify",
                "passed": True,
                "commands": 5,
                "duration_ms": round((time.monotonic() - started) * 1000),
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()

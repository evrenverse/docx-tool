"""Synthetic fixtures for docx-tool."""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
from PIL import Image


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    path = tmp_path / "sample.docx"
    document = Document()
    title = document.add_paragraph(style="Title")
    title.add_run("Quarterly ").bold = True
    title.add_run("Report")
    document.add_paragraph("Owner: Example Team")
    document.add_paragraph("Replace TARGET in this paragraph.")
    document.add_paragraph("Image: [[LOGO]]")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "10"
    document.save(path)
    return path


@pytest.fixture
def controls_docx(tmp_path: Path) -> Path:
    path = tmp_path / "controls.docx"
    document = Document()
    document.add_paragraph("Form fields")

    checkbox = OxmlElement("w:sdt")
    checkbox_properties = OxmlElement("w:sdtPr")
    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), "Approved")
    checkbox_properties.append(alias)
    checkbox_definition = OxmlElement("w14:checkbox")
    checked = OxmlElement("w14:checked")
    checked.set(qn("w14:val"), "0")
    checkbox_definition.append(checked)
    checkbox_properties.append(checkbox_definition)
    checkbox.append(checkbox_properties)
    checkbox_content = OxmlElement("w:sdtContent")
    checkbox_run = OxmlElement("w:r")
    checkbox_text = OxmlElement("w:t")
    checkbox_text.text = "☐ Approved"
    checkbox_run.append(checkbox_text)
    checkbox_content.append(checkbox_run)
    checkbox.append(checkbox_content)

    control = OxmlElement("w:sdt")
    control_properties = OxmlElement("w:sdtPr")
    control_alias = OxmlElement("w:alias")
    control_alias.set(qn("w:val"), "ProjectName")
    control_properties.append(control_alias)
    control.append(control_properties)
    control_content = OxmlElement("w:sdtContent")
    control_run = OxmlElement("w:r")
    control_text = OxmlElement("w:t")
    control_text.text = "Example Project"
    control_run.append(control_text)
    control_content.append(control_run)
    control.append(control_content)

    body = document.element.body
    body.insert(len(body) - 1, checkbox)
    body.insert(len(body) - 1, control)
    document.save(path)
    return path


def _rewrite_document_xml(path: Path, transform) -> None:
    with zipfile.ZipFile(path, "r") as source:
        entries = [(member, source.read(member)) for member in source.infolist()]
    temporary = path.with_suffix(".rewrite")
    parser = etree.XMLParser(resolve_entities=False, no_network=True)
    with zipfile.ZipFile(temporary, "w") as destination:
        for member, data in entries:
            if member.filename == "word/document.xml":
                root = etree.fromstring(data, parser)
                transform(root)
                data = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
            destination.writestr(member, data)
    temporary.replace(path)


@pytest.fixture
def tracked_docx(tmp_path: Path) -> Path:
    path = tmp_path / "tracked.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Before ")
    paragraph.add_run("removed")
    paragraph.add_run("added")
    paragraph.add_run(" after")
    document.save(path)

    def transform(root) -> None:
        paragraph = root.find(f".//{qn('w:p')}")
        runs = list(paragraph.findall(qn("w:r")))
        deleted_run = runs[1]
        inserted_run = runs[2]

        deleted_index = paragraph.index(deleted_run)
        paragraph.remove(deleted_run)
        deleted_text = deleted_run.find(qn("w:t"))
        deleted_text.tag = qn("w:delText")
        deletion = etree.Element(qn("w:del"))
        deletion.set(qn("w:author"), "Reviewer")
        deletion.append(deleted_run)
        paragraph.insert(deleted_index, deletion)

        inserted_index = paragraph.index(inserted_run)
        paragraph.remove(inserted_run)
        insertion = etree.Element(qn("w:ins"))
        insertion.set(qn("w:author"), "Reviewer")
        insertion.append(inserted_run)
        paragraph.insert(inserted_index, insertion)

    _rewrite_document_xml(path, transform)
    return path


@pytest.fixture
def unsupported_revision_docx(tmp_path: Path) -> Path:
    path = tmp_path / "unsupported.docx"
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell"
    document.save(path)

    def transform(root) -> None:
        properties = root.find(f".//{qn('w:tcPr')}")
        properties.append(etree.Element(qn("w:cellMerge")))

    _rewrite_document_xml(path, transform)
    return path


@pytest.fixture
def image_png(tmp_path: Path) -> Path:
    path = tmp_path / "image.png"
    Image.new("RGB", (40, 20), "blue").save(path)
    return path

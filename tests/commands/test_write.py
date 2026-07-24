import json
import os

import pytest
from docx import Document
from typer.testing import CliRunner

from docx_tool.cli import app
from docx_tool.commands.document import extract_checkboxes, extract_content_controls

runner = CliRunner()


def invoke_write(source, output, operations):
    return runner.invoke(
        app,
        ["write", str(source), "-", "--output", str(output), "--json"],
        input=json.dumps(operations),
    )


def test_replace_across_runs_preserves_first_run_format(sample_docx, tmp_path):
    output = tmp_path / "result.docx"
    result = invoke_write(
        sample_docx,
        output,
        [{"type": "replace", "search": "Quarterly Report", "value": "Annual Summary"}],
    )
    assert result.exit_code == 0, result.output
    document = Document(output)
    assert document.paragraphs[0].text == "Annual Summary"
    assert document.paragraphs[0].runs[0].bold is True


def test_replace_all_matches(sample_docx, tmp_path):
    document = Document(sample_docx)
    document.add_paragraph("TARGET and TARGET")
    document.save(sample_docx)
    output = tmp_path / "result.docx"
    result = invoke_write(
        sample_docx,
        output,
        [{"type": "replace", "search": "TARGET", "value": "done"}],
    )
    assert result.exit_code == 0
    assert json.loads(result.output)["applied"][0]["matches"] == 3


def test_replace_does_not_loop_when_replacement_contains_search(sample_docx, tmp_path):
    output = tmp_path / "result.docx"
    result = invoke_write(
        sample_docx,
        output,
        [{"type": "replace", "search": "TARGET", "value": "TARGET-DONE"}],
    )
    assert result.exit_code == 0
    assert Document(output).paragraphs[2].text == "Replace TARGET-DONE in this paragraph."


def test_paragraph_and_table_cell(sample_docx, tmp_path):
    output = tmp_path / "result.docx"
    result = invoke_write(
        sample_docx,
        output,
        [
            {"type": "paragraph", "index": 2, "value": "Owner: New Team"},
            {"type": "table_cell", "table": 1, "row": 2, "column": 2, "value": "42"},
        ],
    )
    assert result.exit_code == 0
    document = Document(output)
    assert document.paragraphs[1].text == "Owner: New Team"
    assert document.tables[0].cell(1, 1).text == "42"


def test_checkbox_and_content_control(controls_docx, tmp_path):
    output = tmp_path / "result.docx"
    result = invoke_write(
        controls_docx,
        output,
        [
            {"type": "checkbox", "id": "Approved", "checked": True},
            {"type": "content_control", "id": "ProjectName", "value": "Public Project"},
        ],
    )
    assert result.exit_code == 0
    document = Document(output)
    assert extract_checkboxes(document)[0]["checked"] is True
    assert extract_content_controls(document)[0]["value"] == "Public Project"


def test_image_marker(sample_docx, image_png, tmp_path):
    output = tmp_path / "result.docx"
    result = invoke_write(
        sample_docx,
        output,
        [
            {
                "type": "image",
                "search": "[[LOGO]]",
                "image": str(image_png),
                "width_inches": 1,
            }
        ],
    )
    assert result.exit_code == 0
    document = Document(output)
    assert len(document.inline_shapes) == 1


def test_failure_is_transactional(sample_docx, tmp_path):
    output = tmp_path / "result.docx"
    result = invoke_write(
        sample_docx,
        output,
        [
            {"type": "paragraph", "index": 2, "value": "would change"},
            {"type": "replace", "search": "missing text", "value": "x"},
        ],
    )
    assert result.exit_code == 1
    assert not output.exists()
    data = json.loads(result.output)
    assert data["written"] is False
    assert data["applied"] == []


def test_in_place_failure_keeps_original(sample_docx):
    before = sample_docx.read_bytes()
    result = runner.invoke(
        app,
        ["write", str(sample_docx), "-", "--json"],
        input=json.dumps([{"type": "unknown"}]),
    )
    assert result.exit_code == 1
    assert sample_docx.read_bytes() == before


@pytest.mark.skipif(os.name != "posix", reason="POSIX mode bits")
def test_in_place_write_preserves_mode(sample_docx):
    sample_docx.chmod(0o640)
    result = runner.invoke(
        app,
        ["write", str(sample_docx), "-"],
        input=json.dumps([{"type": "paragraph", "index": 2, "value": "Updated"}]),
    )
    assert result.exit_code == 0
    assert sample_docx.stat().st_mode & 0o777 == 0o640


def test_write_rejects_non_utf8_changes(tmp_path, sample_docx):
    """A legacy code-page edit plan exits cleanly instead of raising."""
    changes = tmp_path / "cp1252.json"
    changes.write_bytes('[{"type": "replace", "search": "Größe", "value": "x"}]'.encode("cp1252"))
    result = runner.invoke(app, ["write", str(sample_docx), str(changes)])
    assert result.exit_code == 1
    assert "not valid UTF-8" in result.output

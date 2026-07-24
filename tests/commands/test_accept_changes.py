import json
import zipfile

from docx import Document
from typer.testing import CliRunner

from docx_tool.cli import app

runner = CliRunner()


def test_accept_insertions_and_deletions(tracked_docx, tmp_path):
    output = tmp_path / "accepted.docx"
    result = runner.invoke(
        app,
        ["accept-changes", str(tracked_docx), "--output", str(output), "--json"],
    )
    assert result.exit_code == 0, result.output
    data = json.loads(result.output)
    assert data["accepted"] == 2
    assert Document(output).paragraphs[0].text == "Before added after"
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml")
    assert b"<w:ins" not in xml
    assert b"<w:del" not in xml


def test_accept_changes_in_place(tracked_docx):
    result = runner.invoke(app, ["accept-changes", str(tracked_docx)])
    assert result.exit_code == 0
    assert Document(tracked_docx).paragraphs[0].text == "Before added after"


def test_unsupported_change_keeps_original(unsupported_revision_docx, tmp_path):
    before = unsupported_revision_docx.read_bytes()
    output = tmp_path / "result.docx"
    result = runner.invoke(
        app,
        ["accept-changes", str(unsupported_revision_docx), "--output", str(output)],
    )
    assert result.exit_code == 1
    assert "cellMerge" in result.output
    assert not output.exists()
    assert unsupported_revision_docx.read_bytes() == before

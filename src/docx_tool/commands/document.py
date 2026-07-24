"""Read-only document-model helpers based on python-docx."""

from __future__ import annotations

from collections.abc import Iterator
from typing import Any

from docx.oxml.ns import qn


def _control_name(properties) -> str:
    alias = properties.find(qn("w:alias"))
    tag = properties.find(qn("w:tag"))
    if alias is not None:
        return alias.get(qn("w:val"), "")
    if tag is not None:
        return tag.get(qn("w:val"), "")
    return ""


def extract_checkboxes(document) -> list[dict]:
    """Return structured checkbox content-control values."""
    result = []
    for sdt in document.element.iter(qn("w:sdt")):
        properties = sdt.find(qn("w:sdtPr"))
        if properties is None:
            continue
        checkbox = properties.find(qn("w14:checkbox"))
        if checkbox is None:
            continue
        checked_element = checkbox.find(qn("w14:checked"))
        checked = checked_element is not None and checked_element.get(
            qn("w14:val"), "0"
        ).lower() in {"1", "true", "on"}
        content = sdt.find(qn("w:sdtContent"))
        text = (
            "".join(node.text or "" for node in content.iter(qn("w:t")))
            if content is not None
            else ""
        )
        result.append(
            {
                "id": _control_name(properties),
                "text": text,
                "checked": checked,
            }
        )
    return result


def extract_content_controls(document) -> list[dict]:
    """Return named, non-checkbox content controls."""
    result = []
    for sdt in document.element.iter(qn("w:sdt")):
        properties = sdt.find(qn("w:sdtPr"))
        if properties is None or properties.find(qn("w14:checkbox")) is not None:
            continue
        name = _control_name(properties)
        if not name:
            continue
        control_type = "date" if properties.find(qn("w:date")) is not None else "text"
        content = sdt.find(qn("w:sdtContent"))
        text = (
            "".join(node.text or "" for node in content.iter(qn("w:t")))
            if content is not None
            else ""
        )
        result.append({"id": name, "type": control_type, "value": text})
    return result


def extract_revisions(document) -> list[dict]:
    """Return insertion, deletion, and move revision summaries."""
    revisions = []
    for revision_type, element_name, text_name in (
        ("insertion", "w:ins", "w:t"),
        ("deletion", "w:del", "w:delText"),
        ("move_from", "w:moveFrom", "w:delText"),
        ("move_to", "w:moveTo", "w:t"),
    ):
        for element in document.element.iter(qn(element_name)):
            text = "".join(node.text or "" for node in element.iter(qn(text_name)))
            revisions.append(
                {
                    "type": revision_type,
                    "author": element.get(qn("w:author"), ""),
                    "date": element.get(qn("w:date"), ""),
                    "text": text,
                }
            )
    return revisions


def iter_table_cells(document) -> Iterator[tuple[int, int, int, Any]]:
    """Yield unique table cells using one-based indices."""
    for table_index, table in enumerate(document.tables, 1):
        seen = set()
        for row_index, row in enumerate(table.rows, 1):
            for column_index, cell in enumerate(row.cells, 1):
                identity = id(cell._tc)
                if identity in seen:
                    continue
                seen.add(identity)
                yield table_index, row_index, column_index, cell


def iter_editable_paragraphs(document) -> Iterator[tuple[str, Any]]:
    """Yield body and unique table-cell paragraphs with stable locations."""
    for index, paragraph in enumerate(document.paragraphs, 1):
        yield f"P{index}", paragraph
    for table_index, row_index, column_index, cell in iter_table_cells(document):
        for paragraph_index, paragraph in enumerate(cell.paragraphs, 1):
            yield (
                f"T{table_index}:R{row_index}:C{column_index}:P{paragraph_index}",
                paragraph,
            )


def comments_data(document) -> list[dict]:
    """Return simple comment metadata and text."""
    return [
        {
            "id": comment.comment_id,
            "author": comment.author,
            "initials": comment.initials,
            "text": comment.text,
        }
        for comment in document.comments
    ]

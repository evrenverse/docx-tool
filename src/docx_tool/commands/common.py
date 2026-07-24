"""Shared validation, atomic-output, and LibreOffice helpers."""

from __future__ import annotations

import os
import shutil
import subprocess
import tempfile
import zipfile
from collections.abc import Iterator
from contextlib import contextmanager
from pathlib import Path

import typer
from docx import Document

MAX_INPUT_BYTES = 256 * 1024 * 1024
MAX_JSON_INPUT_BYTES = 10 * 1024 * 1024
MAX_BATCH_ITEMS = 10_000


@contextmanager
def atomic_output(path: str | Path) -> Iterator[Path]:
    """Yield a sibling temp path and publish it atomically after success."""
    target = Path(path)
    if not target.parent.is_dir():
        raise ValueError(f"output directory does not exist: {target.parent}")
    original_mode = target.stat().st_mode if target.exists() else None
    fd, temporary_name = tempfile.mkstemp(
        dir=target.parent,
        prefix=f".{target.name}.",
        suffix=".tmp",
    )
    os.close(fd)
    temporary = Path(temporary_name)
    try:
        yield temporary
        sync_fd = os.open(temporary, os.O_RDONLY)
        try:
            os.fsync(sync_fd)
        finally:
            os.close(sync_fd)
        if original_mode is not None:
            os.chmod(temporary, original_mode & 0o777)
        os.replace(temporary, target)
        if os.name == "posix":
            directory_fd = os.open(target.parent, os.O_RDONLY)
            try:
                os.fsync(directory_fd)
            finally:
                os.close(directory_fd)
    finally:
        temporary.unlink(missing_ok=True)


def require_file(path: Path, suffix: str) -> None:
    """Exit with a clean message unless path is an existing regular file."""
    if not path.is_file():
        typer.echo(f"Error: file not found: {path}", err=True)
        raise typer.Exit(code=1)
    if path.suffix.lower() != suffix:
        typer.echo(f"Error: expected a {suffix} file: {path}", err=True)
        raise typer.Exit(code=1)
    size = path.stat().st_size
    if size > MAX_INPUT_BYTES:
        typer.echo(
            f"Error: input is {size} bytes; maximum is {MAX_INPUT_BYTES} bytes",
            err=True,
        )
        raise typer.Exit(code=2)


def ensure_json_size(raw: str) -> None:
    """Reject oversized edit plans before JSON decoding."""
    size = len(raw.encode("utf-8"))
    if size > MAX_JSON_INPUT_BYTES:
        typer.echo(
            f"Error: JSON input is {size} bytes; maximum is {MAX_JSON_INPUT_BYTES} bytes",
            err=True,
        )
        raise typer.Exit(code=2)


def load_document(path: Path):
    """Open a DOCX or exit with a concise diagnostic."""
    require_file(path, ".docx")
    try:
        return Document(str(path))
    except Exception as exc:
        typer.echo(f"Error: cannot open DOCX: {exc}", err=True)
        raise typer.Exit(code=1)


def reject_digitally_signed(path: Path) -> None:
    """Refuse a mutation that would silently invalidate an OOXML signature."""
    try:
        with zipfile.ZipFile(path) as archive:
            signed = any(
                member.filename.lower().startswith("_xmlsignatures/")
                for member in archive.infolist()
            )
    except (OSError, zipfile.BadZipFile):
        return
    if signed:
        typer.echo(
            "Error: digitally signed DOCX files are not rewritten because "
            "the signature would become invalid",
            err=True,
        )
        raise typer.Exit(code=1)


def save_document(document, output: Path) -> None:
    """Save and reopen a document before atomically publishing it."""
    try:
        with atomic_output(output) as temporary:
            document.save(str(temporary))
            Document(str(temporary))
    except typer.Exit:
        raise
    except Exception as exc:
        typer.echo(f"Error: cannot write DOCX: {exc}", err=True)
        raise typer.Exit(code=1)


def find_libreoffice() -> str:
    """Return a LibreOffice executable or exit cleanly."""
    for candidate in ("soffice", "libreoffice"):
        if executable := shutil.which(candidate):
            return executable
    typer.echo(
        "Error: LibreOffice not found; install LibreOffice and ensure "
        "'soffice' or 'libreoffice' is on PATH",
        err=True,
    )
    raise typer.Exit(code=1)


def libreoffice_convert(source: Path, output: Path, target_format: str) -> None:
    """Convert source in an isolated profile and atomically publish output."""
    executable = find_libreoffice()
    if not output.parent.is_dir():
        typer.echo(f"Error: output directory does not exist: {output.parent}", err=True)
        raise typer.Exit(code=1)

    try:
        # Keep the profile path short. LibreOffice uses local sockets whose
        # platform path limit can be exceeded by deeply nested output paths.
        with tempfile.TemporaryDirectory(prefix="docx-tool-lo-") as temporary_dir:
            root = Path(temporary_dir)
            input_dir = root / "input"
            converted_dir = root / "converted"
            profile_dir = root / "profile"
            for directory in (input_dir, converted_dir, profile_dir):
                directory.mkdir(mode=0o700)

            private_source = input_dir / source.name
            shutil.copyfile(source, private_source)
            command = [
                executable,
                f"-env:UserInstallation={profile_dir.resolve().as_uri()}",
                "--headless",
                "--nologo",
                "--nodefault",
                "--norestore",
                "--nofirststartwizard",
                "--convert-to",
                target_format,
                "--outdir",
                str(converted_dir),
                str(private_source),
            ]
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                timeout=120,
                check=False,
            )
            if result.returncode != 0:
                detail = (result.stderr or result.stdout).strip()
                raise RuntimeError(f"LibreOffice conversion failed: {detail}")

            produced = converted_dir / f"{source.stem}.{target_format}"
            if not produced.is_file():
                raise RuntimeError("LibreOffice did not produce the expected output file")
            if target_format == "docx":
                Document(str(produced))
            elif target_format == "pdf" and not produced.read_bytes().startswith(b"%PDF-"):
                raise RuntimeError("LibreOffice produced an invalid PDF")

            with atomic_output(output) as temporary:
                shutil.copyfile(produced, temporary)
    except typer.Exit:
        raise
    except subprocess.TimeoutExpired:
        typer.echo("Error: LibreOffice conversion timed out after 120 seconds", err=True)
        raise typer.Exit(code=1)
    except Exception as exc:
        typer.echo(f"Error: {exc}", err=True)
        raise typer.Exit(code=1)

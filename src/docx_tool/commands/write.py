"""Apply validated, all-or-nothing DOCX edits through python-docx."""

from __future__ import annotations

import json
import sys
from pathlib import Path

import typer
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from docx_tool.commands.common import (
    MAX_BATCH_ITEMS,
    ensure_json_size,
    load_document,
    reject_digitally_signed,
    save_document,
)
from docx_tool.commands.document import _control_name, iter_editable_paragraphs


def _replace_once(paragraph, search: str, replacement: str, start_at: int) -> int | None:
    """Replace one match across adjacent runs, preserving unaffected runs."""
    runs = list(paragraph.runs)
    full_text = "".join(run.text for run in runs)
    start = full_text.find(search, start_at)
    if start < 0:
        return None
    end = start + len(search)

    boundaries = []
    position = 0
    for run in runs:
        boundaries.append((position, position + len(run.text)))
        position += len(run.text)

    start_run = next(
        (index for index, (left, right) in enumerate(boundaries) if left <= start < right),
        None,
    )
    end_run = next(
        (index for index, (left, right) in enumerate(boundaries) if left < end <= right),
        None,
    )
    if start_run is None or end_run is None:
        return None

    start_offset = start - boundaries[start_run][0]
    end_offset = end - boundaries[end_run][0]
    if start_run == end_run:
        text = runs[start_run].text
        runs[start_run].text = text[:start_offset] + replacement + text[end_offset:]
        return start

    runs[start_run].text = runs[start_run].text[:start_offset] + replacement
    for index in range(start_run + 1, end_run):
        runs[index].text = ""
    runs[end_run].text = runs[end_run].text[end_offset:]
    return start


def _replace(document, search: str, replacement: str, first_only: bool) -> int:
    count = 0
    for _, paragraph in iter_editable_paragraphs(document):
        start_at = 0
        while True:
            match = _replace_once(paragraph, search, replacement, start_at)
            if match is None:
                break
            count += 1
            if first_only:
                return count
            start_at = match + len(replacement)
    return count


def _set_paragraph_text(paragraph, value: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def _set_checkbox(document, identifier: str, checked: bool) -> bool:
    for sdt in document.element.iter(qn("w:sdt")):
        properties = sdt.find(qn("w:sdtPr"))
        if properties is None or _control_name(properties) != identifier:
            continue
        checkbox = properties.find(qn("w14:checkbox"))
        if checkbox is None:
            continue
        state = checkbox.find(qn("w14:checked"))
        if state is None:
            state = OxmlElement("w14:checked")
            checkbox.append(state)
        state.set(qn("w14:val"), "1" if checked else "0")
        content = sdt.find(qn("w:sdtContent"))
        if content is not None:
            for text in content.iter(qn("w:t")):
                value = text.text or ""
                if checked:
                    text.text = value.replace("☐", "☒").replace("☑", "☒")
                else:
                    text.text = value.replace("☒", "☐").replace("☑", "☐")
        return True
    return False


def _set_content_control(document, identifier: str, value: str) -> bool:
    for sdt in document.element.iter(qn("w:sdt")):
        properties = sdt.find(qn("w:sdtPr"))
        if properties is None or _control_name(properties) != identifier:
            continue
        if properties.find(qn("w14:checkbox")) is not None:
            continue
        content = sdt.find(qn("w:sdtContent"))
        if content is None:
            return False
        text_nodes = list(content.iter(qn("w:t")))
        if not text_nodes:
            return False
        text_nodes[0].text = value
        for node in text_nodes[1:]:
            node.text = ""
        return True
    return False


def _insert_image(document, search: str, image: Path, width_inches: float) -> bool:
    for _, paragraph in iter_editable_paragraphs(document):
        for run in paragraph.runs:
            if search in run.text:
                before, _, after = run.text.partition(search)
                run.text = before
                run.add_picture(str(image), width=Inches(width_inches))
                if after:
                    run.add_text(after)
                return True
    return False


def _positive_index(value: object, name: str) -> int:
    if isinstance(value, bool) or not isinstance(value, int) or value < 1:
        raise ValueError(f"{name} must be a positive one-based integer")
    return value


def _apply_operation(document, operation: dict) -> dict:
    operation_type = operation.get("type")
    if operation_type == "replace":
        search = operation.get("search")
        replacement = operation.get("value")
        if not isinstance(search, str) or not search:
            raise ValueError("replace.search must be a non-empty string")
        if not isinstance(replacement, str):
            raise ValueError("replace.value must be a string")
        first_only = operation.get("first_only", False)
        if not isinstance(first_only, bool):
            raise ValueError("replace.first_only must be a boolean")
        matches = _replace(document, search, replacement, first_only)
        if matches == 0:
            raise ValueError("search text not found in editable runs")
        return {"type": "replace", "search": search, "value": replacement, "matches": matches}

    if operation_type == "paragraph":
        index = _positive_index(operation.get("index"), "paragraph.index")
        value = operation.get("value")
        if not isinstance(value, str):
            raise ValueError("paragraph.value must be a string")
        if index > len(document.paragraphs):
            raise ValueError(f"paragraph index out of range (1-{len(document.paragraphs)})")
        _set_paragraph_text(document.paragraphs[index - 1], value)
        return {"type": "paragraph", "index": index, "value": value}

    if operation_type == "table_cell":
        table = _positive_index(operation.get("table"), "table_cell.table")
        row = _positive_index(operation.get("row"), "table_cell.row")
        column = _positive_index(operation.get("column", operation.get("col")), "table_cell.column")
        value = operation.get("value")
        if not isinstance(value, str):
            raise ValueError("table_cell.value must be a string")
        if table > len(document.tables):
            raise ValueError(f"table index out of range (1-{len(document.tables)})")
        current = document.tables[table - 1]
        if row > len(current.rows) or column > len(current.columns):
            raise ValueError("table row or column out of range")
        current.cell(row - 1, column - 1).text = value
        return {
            "type": "table_cell",
            "table": table,
            "row": row,
            "column": column,
            "value": value,
        }

    if operation_type == "checkbox":
        identifier = operation.get("id")
        checked = operation.get("checked")
        if not isinstance(identifier, str) or not identifier:
            raise ValueError("checkbox.id must be a non-empty string")
        if not isinstance(checked, bool):
            raise ValueError("checkbox.checked must be a boolean")
        if not _set_checkbox(document, identifier, checked):
            raise ValueError("checkbox id not found")
        return {"type": "checkbox", "id": identifier, "checked": checked}

    if operation_type == "content_control":
        identifier = operation.get("id")
        value = operation.get("value")
        if not isinstance(identifier, str) or not identifier:
            raise ValueError("content_control.id must be a non-empty string")
        if not isinstance(value, str):
            raise ValueError("content_control.value must be a string")
        if not _set_content_control(document, identifier, value):
            raise ValueError("content control id not found or has no text")
        return {"type": "content_control", "id": identifier, "value": value}

    if operation_type == "image":
        search = operation.get("search")
        image_value = operation.get("image")
        width = operation.get("width_inches", 2.0)
        if not isinstance(search, str) or not search:
            raise ValueError("image.search must be a non-empty string")
        if not isinstance(image_value, str):
            raise ValueError("image.image must be a path string")
        image = Path(image_value)
        if not image.is_file():
            raise ValueError(f"image file not found: {image}")
        if isinstance(width, bool) or not isinstance(width, int | float) or width <= 0:
            raise ValueError("image.width_inches must be a positive number")
        if not _insert_image(document, search, image, float(width)):
            raise ValueError("image marker not found inside one editable run")
        return {"type": "image", "search": search, "image": str(image), "width_inches": width}

    raise ValueError(f"unknown operation type: {operation_type}")


def write(
    file: Path = typer.Argument(..., help="Path to the DOCX file."),
    changes: str = typer.Argument(..., help="JSON file with operations, or '-' for stdin."),
    output: Path | None = typer.Option(None, "--output", "-o", help="Output DOCX path."),
    output_json: bool = typer.Option(False, "--json", help="Output a structured result."),
) -> None:
    """Apply all operations or write nothing when any operation fails."""
    document = load_document(file)
    reject_digitally_signed(file)
    if changes == "-":
        raw = sys.stdin.read()
    else:
        changes_path = Path(changes)
        if not changes_path.is_file():
            typer.echo(f"Error: changes file not found: {changes}", err=True)
            raise typer.Exit(code=1)
        raw = changes_path.read_text(encoding="utf-8")
    ensure_json_size(raw)

    try:
        operations = json.loads(raw)
    except json.JSONDecodeError as exc:
        typer.echo(f"Error: invalid JSON: {exc}", err=True)
        raise typer.Exit(code=1)
    if not isinstance(operations, list):
        typer.echo("Error: JSON must be an array of operation objects", err=True)
        raise typer.Exit(code=1)
    if len(operations) > MAX_BATCH_ITEMS:
        typer.echo(f"Error: batch exceeds {MAX_BATCH_ITEMS} operations", err=True)
        raise typer.Exit(code=2)

    applied = []
    failures = []
    for index, operation in enumerate(operations):
        if not isinstance(operation, dict):
            failures.append({"index": index, "reason": "operation must be an object"})
            continue
        try:
            applied.append(_apply_operation(document, operation))
        except (TypeError, ValueError) as exc:
            failures.append(
                {
                    "index": index,
                    "type": operation.get("type"),
                    "reason": str(exc),
                }
            )

    output_path = output or file
    if failures:
        if output_json:
            typer.echo(
                json.dumps(
                    {
                        "file": output_path.name,
                        "applied": [],
                        "failed": failures,
                        "written": False,
                    },
                    indent=2,
                    ensure_ascii=False,
                )
            )
        else:
            for failure in failures:
                typer.echo(
                    f"Error: operation {failure['index']}: {failure['reason']}",
                    err=True,
                )
        raise typer.Exit(code=1)

    save_document(document, output_path)
    result = {
        "file": output_path.name,
        "applied": applied,
        "failed": [],
        "written": True,
    }
    if output_json:
        typer.echo(json.dumps(result, indent=2, ensure_ascii=False))
    else:
        typer.echo(f"Written: {len(applied)} operations to {output_path.name}")
